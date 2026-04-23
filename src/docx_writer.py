import os
import re
import sys

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Inches, RGBColor

# Ensure project root is importable when this module is run directly
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from config import (
    CANDIDATE_NAME, CANDIDATE_LINKEDIN_URL,
    CANDIDATE_EMAIL, CANDIDATE_PHONE, CANDIDATE_LOCATION,
)

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUTPUT_DIR = os.path.join(BASE_DIR, "output")

# ---------------------------------------------------------------------------
# Date-line detection (for right-aligned dates)
# ---------------------------------------------------------------------------

_MONTH = r'(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|Sep(?:t(?:ember)?)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)'
_SEASON = r'(?:Spring|Summer|Fall|Autumn|Winter)'
_DATE_RE = re.compile(
    r'^(.+?)\s*\|\s*'
    r'(' + _MONTH + r'\s+\d{4}\s*[–—\-]\s*(?:Present|Current|Ongoing|' + _MONTH + r'\s+\d{4})'
    r'|' + _SEASON + r'\s+\d{4}\s*[–—\-]\s*(?:Present|Current|Ongoing|' + _SEASON + r'\s+\d{4}|' + _MONTH + r'\s+\d{4})'
    r'|\d{4}(?:\s*[–—\-]\s*(?:Present|Current|Ongoing|\d{4}))?'
    r'.*)$'
)


def _parse_date_line(line: str):
    """Return (left_text, date_text) if *line* has a pipe-separated date, else None."""
    m = _DATE_RE.match(line)
    if m:
        return m.group(1).strip(), m.group(2).strip()
    return None

# ---------------------------------------------------------------------------
# Sector detection
# ---------------------------------------------------------------------------

_CONSULTING_KEYWORDS = [
    "consulting", "consultant", "advisory", "management consulting",
    "mckinsey", "bcg", "bain", "deloitte", "pwc", "accenture", "kpmg",
    "strategy&", "oliver wyman", "monitor", "analysis group",
]
_CLEAN_ENERGY_KEYWORDS = [
    "clean energy", "renewable", "solar", "wind", "climate", "net zero",
    "carbon", "energy transition", "green finance", "sustainable finance",
    "project finance", "infrastructure fund", "energy storage", "battery",
    "decarbonization", "decarbonisation", "cleantech", "power purchase",
    "bloomberg nef", "bnef", "private equity", "infrastructure",
]
_NGO_KEYWORDS = [
    "foundation", "philanthropy", "philanthropic", "nonprofit", "non-profit",
    "ngo", "charity", "grantmaking", "impact investing", "social impact",
    "development", "international development", "civil society", "advocacy",
    "giin", "ifc", "world bank", "usaid", "undp", "oecd", "idb", "iadb",
    "effective altruism", "evaluation", "monitoring", "learning", "mel",
]


def detect_sector(text: str) -> str:
    """Infer sector from job description / role text.

    Returns one of: 'consulting', 'clean_energy', 'ngo_philanthropy'.
    Defaults to 'ngo_philanthropy' when ambiguous.
    """
    lower = text.lower()

    consulting_hits = sum(1 for kw in _CONSULTING_KEYWORDS if kw in lower)
    clean_energy_hits = sum(1 for kw in _CLEAN_ENERGY_KEYWORDS if kw in lower)
    ngo_hits = sum(1 for kw in _NGO_KEYWORDS if kw in lower)

    scores = {
        "consulting": consulting_hits,
        "clean_energy": clean_energy_hits,
        "ngo_philanthropy": ngo_hits,
    }
    best = max(scores, key=lambda k: scores[k])
    # If all scores are 0 or tied at 0, fall back to ngo_philanthropy
    if scores[best] == 0:
        return "ngo_philanthropy"
    return best


# ---------------------------------------------------------------------------
# Per-sector formatting settings
# ---------------------------------------------------------------------------

_SECTOR_FORMATS = {
    "consulting": {
        "body_pt": 10.5,
        "header_name_pt": 13.0,
        "section_heading_pt": 10.5,
        "top_margin": 0.50,
        "bottom_margin": 0.50,
        "left_margin": 0.70,
        "right_margin": 0.70,
        "line_spacing_pt": 11.5,
        "space_after_pt": 1.0,
        "bullet_indent": 0.20,
        "heading_space_before_pt": 4.0,
    },
    "clean_energy": {
        "body_pt": 10.5,
        "header_name_pt": 13.0,
        "section_heading_pt": 10.5,
        "top_margin": 0.50,
        "bottom_margin": 0.50,
        "left_margin": 0.75,
        "right_margin": 0.75,
        "line_spacing_pt": 12.0,
        "space_after_pt": 1.5,
        "bullet_indent": 0.25,
        "heading_space_before_pt": 4.0,
    },
    "ngo_philanthropy": {
        "body_pt": 10.5,
        "header_name_pt": 13.0,
        "section_heading_pt": 10.5,
        "top_margin": 0.50,
        "bottom_margin": 0.50,
        "left_margin": 0.75,
        "right_margin": 0.75,
        "line_spacing_pt": 12.0,
        "space_after_pt": 1.5,
        "bullet_indent": 0.25,
        "heading_space_before_pt": 4.0,
    },
}


# ---------------------------------------------------------------------------
# Internal helpers
# ---------------------------------------------------------------------------

def _slugify(text: str) -> str:
    return text.strip().replace(" ", "_").replace("/", "-")


def _apply_normal_style(doc: Document, fmt: dict) -> None:
    style = doc.styles["Normal"]
    style.font.name = fmt.get("font_family", "Calibri")
    style.font.size = Pt(fmt["body_pt"])
    style.paragraph_format.space_after = Pt(fmt["space_after_pt"])
    style.paragraph_format.line_spacing = Pt(fmt["line_spacing_pt"])


def _apply_margins(doc: Document, fmt: dict) -> None:
    for section in doc.sections:
        section.top_margin = Inches(fmt["top_margin"])
        section.bottom_margin = Inches(fmt["bottom_margin"])
        section.left_margin = Inches(fmt["left_margin"])
        section.right_margin = Inches(fmt["right_margin"])


def _add_hyperlink(paragraph, text: str, url: str) -> None:
    """Insert a visible hyperlink into *paragraph* using OOXML."""
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run_elem = OxmlElement("w:r")

    # Run properties: blue, underlined
    rpr = OxmlElement("w:rPr")

    color_elem = OxmlElement("w:color")
    color_elem.set(qn("w:val"), "0563C1")
    rpr.append(color_elem)

    u_elem = OxmlElement("w:u")
    u_elem.set(qn("w:val"), "single")
    rpr.append(u_elem)

    run_elem.append(rpr)

    t = OxmlElement("w:t")
    t.text = text
    run_elem.append(t)

    hyperlink.append(run_elem)
    paragraph._p.append(hyperlink)


def _add_name_contact_block(doc: Document, fmt: dict) -> None:
    """Add the candidate's name and contact info as the first body paragraphs.

    Name on its own centered line (bold, colored, large).
    Contact info on a second centered line (body font) with LinkedIn hyperlink.
    This avoids the name being cut off when sharing a line with long contact text.
    """
    font_name = fmt.get("font_family", "Calibri")
    name_rgb = _hex_to_rgb(fmt.get("name_color")) or _hex_to_rgb(fmt.get("accent_color"))
    contact_style = fmt.get("contact_style", "pipe_separated")
    separator = " | " if contact_style == "pipe_separated" else " · "

    # --- Line 1: Candidate name (centered, bold, colored) ---
    name_para = doc.add_paragraph()
    name_para.alignment = 1  # CENTER
    name_para.paragraph_format.space_after = Pt(1)
    name_para.paragraph_format.space_before = Pt(0)

    name_run = name_para.add_run(CANDIDATE_NAME)
    name_run.bold = True
    name_run.font.name = font_name
    name_run.font.size = Pt(fmt["header_name_pt"])
    if name_rgb:
        name_run.font.color.rgb = name_rgb

    # --- Line 2: Contact info + LinkedIn (centered, body font) ---
    contact_para = doc.add_paragraph()
    contact_para.alignment = 1  # CENTER
    contact_para.paragraph_format.space_after = Pt(2)
    contact_para.paragraph_format.space_before = Pt(0)

    contact_parts = [CANDIDATE_LOCATION, CANDIDATE_PHONE, CANDIDATE_EMAIL]
    contact_text = separator.join(contact_parts) + separator
    contact_run = contact_para.add_run(contact_text)
    contact_run.font.name = font_name
    contact_run.font.size = Pt(fmt["body_pt"])

    # LinkedIn as clickable hyperlink
    _add_hyperlink(contact_para, "LinkedIn", CANDIDATE_LINKEDIN_URL)


def _hex_to_rgb(hex_color: str) -> RGBColor | None:
    """Convert '#2B579A' to RGBColor, returning None on invalid input."""
    if not hex_color:
        return None
    hex_color = hex_color.lstrip("#")
    if len(hex_color) != 6:
        return None
    try:
        return RGBColor(
            int(hex_color[0:2], 16),
            int(hex_color[2:4], 16),
            int(hex_color[4:6], 16),
        )
    except ValueError:
        return None


def _add_heading_underline(paragraph, hex_color: str) -> None:
    """Add a colored bottom border (underline) to a paragraph using OOXML."""
    color = hex_color.lstrip("#") if hex_color else "333333"
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")        # half-point size
    bottom.set(qn("w:space"), "1")     # space between text and border
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_section_divider(doc: Document, fmt: dict) -> None:
    """Add a thin horizontal divider between sections if configured."""
    style = fmt.get("divider_style", "none")
    if style == "none":
        return
    color = fmt.get("accent_color", "#333333").lstrip("#")
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single" if style == "solid_line" else "double")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_line_with_right_date(
    doc: Document, left: str, date: str, fmt: dict, bold: bool = False
) -> None:
    """Add a paragraph with left-aligned text and right-aligned date via tab stop."""
    p = doc.add_paragraph()
    font_name = fmt.get("font_family", "Calibri")

    run_left = p.add_run(left)
    run_left.font.name = font_name
    run_left.font.size = Pt(fmt["body_pt"])
    run_left.bold = bold

    tab_run = p.add_run("\t")
    tab_run.font.name = font_name
    tab_run.font.size = Pt(fmt["body_pt"])

    run_date = p.add_run(date)
    run_date.font.name = font_name
    run_date.font.size = Pt(fmt["body_pt"])
    run_date.bold = bold

    # Calculate right tab position based on page width minus margins
    usable_width_inches = 8.5 - fmt.get("left_margin", 1.0) - fmt.get("right_margin", 1.0)
    tab_pos_twips = str(int(usable_width_inches * 1440))

    pPr = p._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:pos"), tab_pos_twips)
    tabs.append(tab)
    pPr.append(tabs)

    p.paragraph_format.space_after = Pt(fmt["space_after_pt"])


def _is_contact_line(line: str) -> bool:
    """True if the line looks like contact info (email, phone, location separated by pipes/bullets)."""
    # Heuristic: contains @ or phone pattern AND has pipe/bullet separators
    has_email = "@" in line
    has_phone = bool(re.search(r'\+?\d[\d\s\-().]{7,}', line))
    has_separators = "|" in line or "·" in line or "•" in line
    return has_separators and (has_email or has_phone)


def _is_section_heading(line: str) -> bool:
    """True for short all-caps lines, lines ending with ':', or short Title Case standalone lines."""
    stripped = line.strip()
    if not stripped or len(stripped) > 60:
        return False
    # ALL CAPS (primary contract format)
    if stripped.isupper():
        return True
    # Ends with colon
    if stripped.endswith(":") and stripped[0].isupper():
        return True
    # Title Case fallback: all words capitalized, 1-4 words, no date/bullet markers
    words = stripped.split()
    if 1 <= len(words) <= 4 and all(w[0].isupper() for w in words if w[0].isalpha()):
        if not any(c in stripped for c in ["|", "•", "-", "–", "@", "(", ","]):
            return True
    return False


def _is_name_line(line: str) -> bool:
    """True if the line matches the candidate's name (first content line)."""
    return line.strip().lower() == CANDIDATE_NAME.lower()


def _write_content(doc: Document, content: str, fmt: dict) -> None:
    first_non_empty = True
    accent_rgb = _hex_to_rgb(fmt.get("accent_color"))
    name_rgb = _hex_to_rgb(fmt.get("name_color")) or accent_rgb
    use_underline = fmt.get("heading_underline", True)  # default True per reference
    font_name = fmt.get("font_family", "Calibri")
    contact_style = fmt.get("contact_style", "pipe_separated")
    # 3 = WD_ALIGN_PARAGRAPH.JUSTIFY
    body_align = 3

    for line in content.splitlines():
        stripped = line.rstrip()

        if not stripped:
            doc.add_paragraph()
            continue

        # Candidate name + contact line (first non-empty line)
        # In the reference resume, name and contact are on a single justified line
        if first_non_empty and _is_name_line(stripped):
            first_non_empty = False
            # Name is rendered by _add_name_contact_block, skip here
            continue

        if first_non_empty and _is_contact_line(stripped):
            first_non_empty = False
            # Contact rendered by _add_name_contact_block, skip here
            continue

        first_non_empty = False

        # Contact info line (if it appears later, not first line)
        if _is_contact_line(stripped):
            separator = " | " if contact_style == "pipe_separated" else " · "
            parts = re.split(r'\s*[|·•]\s*', stripped)
            formatted = separator.join(p.strip() for p in parts if p.strip())
            p = doc.add_paragraph()
            run = p.add_run(formatted)
            run.font.name = font_name
            run.font.size = Pt(fmt["body_pt"])
            p.alignment = body_align
            p.paragraph_format.space_after = Pt(2)
            continue

        if _is_section_heading(stripped):
            p = doc.add_paragraph()
            run = p.add_run(stripped)
            run.bold = True
            run.font.name = font_name
            run.font.size = Pt(fmt["section_heading_pt"])
            if accent_rgb:
                run.font.color.rgb = accent_rgb
            p.paragraph_format.space_before = Pt(fmt["heading_space_before_pt"])
            p.paragraph_format.space_after = Pt(2)
            if use_underline:
                _add_heading_underline(p, fmt.get("accent_color", "#2B579A"))

        elif stripped.startswith("• ") or stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(stripped[2:])
            run.font.name = font_name
            run.font.size = Pt(fmt["body_pt"])
            p.paragraph_format.left_indent = Inches(fmt["bullet_indent"])
            p.paragraph_format.space_after = Pt(fmt["space_after_pt"])
            p.alignment = body_align  # justified like reference

        else:
            # Check for date-line pattern (e.g. "Company | Jan 2020 – Present")
            date_parts = _parse_date_line(stripped)
            if date_parts:
                _add_line_with_right_date(
                    doc, date_parts[0], date_parts[1], fmt, bold=True
                )
            else:
                p = doc.add_paragraph()
                run = p.add_run(stripped)
                run.font.name = font_name
                run.font.size = Pt(fmt["body_pt"])
                p.paragraph_format.space_after = Pt(fmt["space_after_pt"])
                p.alignment = body_align  # justified like reference


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def _build_fmt(sector: str, layout: dict | None) -> dict:
    """Build the format dict by merging sector defaults with layout overrides."""
    fmt = dict(_SECTOR_FORMATS.get(sector, _SECTOR_FORMATS["ngo_philanthropy"]))
    if layout:
        for key in (
            "body_pt", "header_name_pt", "section_heading_pt",
            "top_margin", "bottom_margin", "left_margin", "right_margin",
            "line_spacing_pt", "space_after_pt", "bullet_indent",
            "heading_space_before_pt",
        ):
            if key in layout and isinstance(layout[key], (int, float)):
                fmt[key] = layout[key]
        for key in (
            "accent_color", "heading_underline", "name_color",
            "divider_style", "contact_style", "font_family",
        ):
            if key in layout:
                fmt[key] = layout[key]
    return fmt


def write_docx(
    content: str,
    company: str,
    role: str,
    doc_type: str,
    sector: str = "ngo_philanthropy",
    layout: dict | None = None,
) -> dict:
    """Save content as a .docx file to output/[company]_[role]/[doc_type].docx.

    Args:
        sector: one of 'consulting', 'clean_energy', 'ngo_philanthropy'.
                Used as fallback when layout is not provided.
        layout: A LayoutSpec dict from the Layout Designer agent.
                When provided, overrides the sector-based format defaults.

    Returns a dict with:
        - path: absolute path of the saved file
        - estimated_pages: int from estimate_page_count
        - page_target: int or None (from layout)
    """
    fmt = _build_fmt(sector, layout)

    folder = os.path.join(OUTPUT_DIR, f"{_slugify(company)}_{_slugify(role)}")
    os.makedirs(folder, exist_ok=True)

    filepath = os.path.join(folder, f"{doc_type}.docx")

    doc = Document()
    _apply_normal_style(doc, fmt)
    _apply_margins(doc, fmt)

    # Remove the default empty paragraph Word inserts on new Document()
    for p in list(doc.paragraphs):
        p._element.getparent().remove(p._element)

    _add_name_contact_block(doc, fmt)
    _write_content(doc, content.strip(), fmt)

    doc.save(filepath)

    est_pages = estimate_page_count(content, fmt)
    page_target = layout.get("page_count_target") if layout else None

    return {
        "path": filepath,
        "estimated_pages": est_pages,
        "page_target": page_target,
    }


def estimate_page_count(content: str, fmt: dict) -> int:
    """Rough estimate of how many pages *content* will take with *fmt*.

    Uses a simple line-height model: each line takes ``line_spacing_pt``
    points, section headings add ``heading_space_before_pt`` extra.
    Available height = (11 − top_margin − bottom_margin) × 72 pts.
    """
    available_height_pt = (
        (11.0 - fmt.get("top_margin", 1.0) - fmt.get("bottom_margin", 1.0))
        * 72.0
    )
    line_height = fmt.get("line_spacing_pt", 13.0)
    heading_extra = fmt.get("heading_space_before_pt", 6.0)

    total_height = 0.0
    for line in content.splitlines():
        stripped = line.rstrip()
        if not stripped:
            total_height += line_height * 0.5  # blank lines are shorter
        elif _is_section_heading(stripped):
            total_height += line_height + heading_extra
        else:
            total_height += line_height

    pages = max(1, int(total_height / available_height_pt + 0.99))
    return pages


# ---------------------------------------------------------------------------
# ATS Validation
# ---------------------------------------------------------------------------

def ats_check(filepath: str) -> list[dict]:
    """Check a generated DOCX for common ATS compatibility issues.

    Returns a list of dicts: {"check": str, "status": "pass"|"warn"|"fail", "detail": str}
    """
    from lxml import etree

    doc = Document(filepath)
    results = []

    # 1. No tables
    if doc.tables:
        results.append({"check": "No tables", "status": "fail",
                        "detail": f"Found {len(doc.tables)} table(s). ATS may not parse table content."})
    else:
        results.append({"check": "No tables", "status": "pass", "detail": "No tables found."})

    # 2. No images
    has_images = any("image" in rel.reltype for rel in doc.part.rels.values())
    if has_images:
        results.append({"check": "No images", "status": "fail",
                        "detail": "Document contains images. ATS cannot read image content."})
    else:
        results.append({"check": "No images", "status": "pass", "detail": "No images found."})

    # 3. No text boxes / shapes
    body_xml = etree.tostring(doc.element.body)
    has_textbox = b'txbxContent' in body_xml or b'wsp:' in body_xml
    if has_textbox:
        results.append({"check": "No text boxes", "status": "fail",
                        "detail": "Document contains text boxes. ATS often skips these."})
    else:
        results.append({"check": "No text boxes", "status": "pass", "detail": "No text boxes found."})

    # 4. No content in Word headers/footers
    header_text = ""
    for section in doc.sections:
        if section.header and section.header.paragraphs:
            for p in section.header.paragraphs:
                header_text += p.text.strip()
        if section.footer and section.footer.paragraphs:
            for p in section.footer.paragraphs:
                header_text += p.text.strip()
    if header_text:
        results.append({"check": "No header/footer content", "status": "warn",
                        "detail": f"Found text in header/footer: '{header_text[:80]}'. Many ATS skip this."})
    else:
        results.append({"check": "No header/footer content", "status": "pass",
                        "detail": "Headers and footers are empty."})

    # 5. Font consistency
    fonts_used = set()
    for para in doc.paragraphs:
        for run in para.runs:
            if run.font.name:
                fonts_used.add(run.font.name)
    if len(fonts_used) > 2:
        results.append({"check": "Font consistency", "status": "warn",
                        "detail": f"Found {len(fonts_used)} fonts: {', '.join(sorted(fonts_used))}. Consider limiting to 1-2."})
    else:
        results.append({"check": "Font consistency", "status": "pass",
                        "detail": f"Fonts used: {', '.join(sorted(fonts_used)) or 'default'}."})

    # 6. Sufficient content
    text_paras = [p for p in doc.paragraphs if p.text.strip()]
    if len(text_paras) < 5:
        results.append({"check": "Sufficient content", "status": "warn",
                        "detail": f"Only {len(text_paras)} non-empty paragraphs. Content may not have rendered."})
    else:
        results.append({"check": "Sufficient content", "status": "pass",
                        "detail": f"{len(text_paras)} content paragraphs."})

    return results


# ---------------------------------------------------------------------------
# DOCX Verification Report
# ---------------------------------------------------------------------------

def verify_docx(filepath: str) -> dict:
    """Extract a formatting report from a generated DOCX for programmatic verification.

    Returns a dict with formatting metadata and any detected warnings.
    """
    doc = Document(filepath)
    report = {"warnings": []}

    # First paragraph text
    first_text = ""
    for p in doc.paragraphs:
        if p.text.strip():
            first_text = p.text.strip()
            break
    report["first_paragraph"] = first_text
    report["has_name_block"] = CANDIDATE_NAME.lower() in first_text.lower() if first_text else False
    if not report["has_name_block"]:
        report["warnings"].append(f"First paragraph does not contain candidate name '{CANDIDATE_NAME}'.")

    # Fonts and sizes
    fonts = set()
    sizes = set()
    for para in doc.paragraphs:
        for run in para.runs:
            if run.font.name:
                fonts.add(run.font.name)
            if run.font.size:
                sizes.add(run.font.size.pt)
    report["fonts_used"] = sorted(fonts)
    report["font_sizes_pt"] = sorted(sizes)

    # Margins (from first section)
    if doc.sections:
        sec = doc.sections[0]
        report["margins"] = {
            "top": round(sec.top_margin / 914400, 2),
            "bottom": round(sec.bottom_margin / 914400, 2),
            "left": round(sec.left_margin / 914400, 2),
            "right": round(sec.right_margin / 914400, 2),
        }
    else:
        report["margins"] = {}
        report["warnings"].append("No sections found in document.")

    # Paragraph and heading counts
    report["paragraph_count"] = len([p for p in doc.paragraphs if p.text.strip()])
    heading_count = 0
    for p in doc.paragraphs:
        for run in p.runs:
            if run.bold and run.font.color and run.font.color.rgb:
                heading_count += 1
                break
    report["heading_count"] = heading_count

    # Estimated pages from doc margins
    if report.get("margins"):
        m = report["margins"]
        available_height = (11.0 - m.get("top", 1.0) - m.get("bottom", 1.0)) * 72
        avg_size = sum(sizes) / len(sizes) if sizes else 11.0
        line_height = avg_size * 1.2
        lines_per_page = available_height / line_height if line_height else 40
        report["estimated_pages"] = max(1, int(report["paragraph_count"] / lines_per_page + 0.99))
    else:
        report["estimated_pages"] = 1

    return report
